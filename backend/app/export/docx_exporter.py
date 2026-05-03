"""DOCX export for persisted quizzes — two-section format.

Section 1: Quiz cards for participants (no answers).
Section 2: Answer key table (page 2).
"""

from __future__ import annotations

import random
import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor

from backend.app.domain.errors import DomainValidationError
from backend.app.domain.models import Option
from backend.app.domain.models import Question
from backend.app.domain.models import Quiz
from backend.app.export.base import ExportedQuizFile

_GREEN  = RGBColor(0x00, 0x80, 0x00)
_PURPLE = RGBColor(0x66, 0x22, 0xCC)
_BLACK  = RGBColor(0x00, 0x00, 0x00)
_LABELS = ["A", "B", "C", "D"]

_NUM_W_TWIPS = 500
_ANS_W_TWIPS = 2800
_TBL_W_TWIPS = 9360


def _set_table_width(tbl, width_twips: int) -> None:
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)


def _set_col_width(cell, width_twips: int) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _question_para(doc: Document, text: str, index: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{index}. {text}")
    run.bold = True
    run.font.size = Pt(12)


def _option_para(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    r_label = p.add_run(f"{label}.  ")
    r_label.bold = True
    r_label.font.color.rgb = _PURPLE
    p.add_run(text)


def _blank_answer_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Ответ:  " + "_" * 40)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


class QuizDocxExporter:
    """Export persisted quizzes into two-section DOCX files.

    Section 1 — quiz cards for participants (no answers).
    Section 2 — answer key table, separated by a page break.
    """

    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def export(self, quiz: Quiz) -> ExportedQuizFile:
        """Render one quiz into a two-section DOCX file."""

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Inches(0.9)
            section.bottom_margin = Inches(0.9)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)

        h = doc.add_heading(quiz.title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.runs[0].font.color.rgb = _BLACK

        for index, question in enumerate(quiz.questions, start=1):
            self._render_question(doc, question, index)

        _add_page_break(doc)
        self._render_answer_key(doc, quiz.questions)

        output = BytesIO()
        doc.save(output)
        return ExportedQuizFile(
            filename=f"{quiz.quiz_id}.docx",
            media_type=self.media_type,
            content_bytes=output.getvalue(),
        )

    def _render_question(self, doc: Document, question: Question, index: int) -> None:
        qt = question.question_type
        if qt in {"single_choice", "true_false"}:
            self._render_choice(doc, question, index)
        elif qt == "fill_blank":
            self._render_fill_blank(doc, question, index)
        elif qt == "short_answer":
            self._render_short_answer(doc, question, index)
        elif qt == "matching":
            self._render_matching(doc, question, index)
        else:
            self._render_choice(doc, question, index)

    def _render_choice(self, doc: Document, question: Question, index: int) -> None:
        _question_para(doc, question.prompt, index)
        if question.question_type == "true_false":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(2)
            for opt in question.options[:2]:
                r = p.add_run(f"☐  {opt.text}     ")
                r.font.size = Pt(12)
        else:
            for i, option in enumerate(question.options[:4]):
                _option_para(doc, _LABELS[i], option.text)

    def _render_fill_blank(self, doc: Document, question: Question, index: int) -> None:
        prompt = re.sub(r"_{2,}", "___________", question.prompt)
        if "___________" not in prompt:
            prompt += " ___________"
        _question_para(doc, prompt, index)

    def _render_short_answer(self, doc: Document, question: Question, index: int) -> None:
        _question_para(doc, question.prompt, index)
        _blank_answer_line(doc)

    def _render_matching(self, doc: Document, question: Question, index: int) -> None:
        _question_para(doc, question.prompt, index)
        pairs  = list(question.matching_pairs)
        lefts  = [p.left  for p in pairs]
        rights = [p.right for p in pairs]
        shuffled = rights[:]
        random.shuffle(shuffled)

        tbl = doc.add_table(rows=len(pairs) + 1, cols=2)
        tbl.style = "Table Grid"
        tbl.autofit = True

        hdr_l, hdr_r = tbl.rows[0].cells
        for cell, txt in ((hdr_l, "Понятие"), (hdr_r, "Определение")):
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.bold = True
            run.underline = True
            run.font.size = Pt(11)

        for row_i, (ltext, rtext) in enumerate(zip(lefts, shuffled), start=1):
            cl, cr = tbl.rows[row_i].cells
            for cell, txt, lbl in ((cl, ltext, _LABELS[row_i - 1]), (cr, rtext, "")):
                p = cell.paragraphs[0]
                if lbl:
                    rl = p.add_run(f"{lbl}.  ")
                    rl.bold = True
                    rl.font.color.rgb = _PURPLE
                rt = p.add_run(txt)
                rt.font.size = Pt(11)

        doc.add_paragraph()

    def _render_answer_key(self, doc: Document, questions: tuple[Question, ...]) -> None:
        h = doc.add_heading("Ответы", level=1)
        run = h.runs[0]
        run.bold = True
        run.font.color.rgb = _BLACK

        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        tbl.autofit = False
        _set_table_width(tbl, _TBL_W_TWIPS)
        q_w = _TBL_W_TWIPS - _NUM_W_TWIPS - _ANS_W_TWIPS

        hdr = tbl.rows[0].cells
        for cell, txt, w, center in zip(
            hdr,
            ("№", "Вопрос", "Правильный ответ"),
            (_NUM_W_TWIPS, q_w, _ANS_W_TWIPS),
            (True, False, False),
        ):
            _set_col_width(cell, w)
            p = cell.paragraphs[0]
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(txt)
            run.bold = True
            run.underline = True
            run.font.size = Pt(11)

        for i, question in enumerate(questions, start=1):
            row = tbl.add_row().cells
            for cell, w in zip(row, (_NUM_W_TWIPS, q_w, _ANS_W_TWIPS)):
                _set_col_width(cell, w)

            num_p = row[0].paragraphs[0]
            num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            num_p.add_run(str(i)).font.size = Pt(11)

            prompt = question.prompt
            short  = prompt[:70] + "…" if len(prompt) > 70 else prompt
            row[1].paragraphs[0].add_run(short).font.size = Pt(10)

            ans_para = row[2].paragraphs[0]
            if question.question_type == "matching":
                for j, pair in enumerate(question.matching_pairs):
                    p = ans_para if j == 0 else row[2].add_paragraph()
                    run = p.add_run(f"{pair.left} → {pair.right}")
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = _GREEN
            else:
                ans_text = self._answer_text(question)
                run = ans_para.add_run(ans_text)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = _GREEN

    def _answer_text(self, question: Question) -> str:
        qt = question.question_type
        if qt in {"fill_blank", "short_answer"}:
            return question.correct_answer or ""
        idx = self._resolve_correct_option_index(question)
        label = _LABELS[idx] if idx < len(_LABELS) else str(idx + 1)
        text  = question.options[idx].text
        return f"{label}. {text}"

    @staticmethod
    def _resolve_correct_option_index(question: Question) -> int:
        if (
            question.correct_option_index is None
            or question.correct_option_index < 0
            or question.correct_option_index >= len(question.options)
        ):
            raise DomainValidationError(
                f"question '{question.question_id}' has invalid correct_option_index for DOCX export"
            )
        return question.correct_option_index

    @staticmethod
    def _resolve_correct_option(question: Question) -> Option:
        if (
            question.correct_option_index is None
            or question.correct_option_index < 0
            or question.correct_option_index >= len(question.options)
        ):
            raise DomainValidationError(
                f"question '{question.question_id}' has invalid correct_option_index for DOCX export"
            )
        return question.options[question.correct_option_index]
