from __future__ import annotations

from dataclasses import replace
from io import BytesIO

import pytest
from docx import Document

from backend.app.domain.errors import DomainValidationError
from backend.app.domain.models import Explanation
from backend.app.domain.models import Option
from backend.app.domain.models import Question
from backend.app.domain.models import Quiz
from backend.app.export.docx_exporter import QuizDocxExporter
from backend.app.export.registry import DEFAULT_QUIZ_EXPORT_REGISTRY


def build_quiz() -> Quiz:
    return Quiz(
        quiz_id="quiz-ru-1",
        document_id="doc-ru-1",
        title="Тренировочный квиз по географии",
        version=3,
        last_edited_at="2026-04-22T12:00:00.000000Z",
        questions=(
            Question(
                question_id="question-1",
                prompt="Какой город является столицей России?",
                options=(
                    Option(option_id="option-1", text="Москва"),
                    Option(option_id="option-2", text="Казань"),
                ),
                correct_option_index=0,
                explanation=Explanation(text="Москва является столицей России."),
            ),
            Question(
                question_id="question-2",
                prompt="Какая река протекает через Санкт-Петербург?",
                options=(
                    Option(option_id="option-1", text="Волга"),
                    Option(option_id="option-2", text="Нева"),
                ),
                correct_option_index=1,
                explanation=None,
            ),
        ),
    )


def read_docx_text(content_bytes: bytes) -> str:
    document = Document(BytesIO(content_bytes))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    parts.append(para.text)
    return "\n".join(parts)


def test_docx_exporter_builds_openable_quiz_document_with_cyrillic_content() -> None:
    exporter = QuizDocxExporter()

    exported_file = exporter.export(build_quiz())
    document_text = read_docx_text(exported_file.content_bytes)

    assert exported_file.filename == "quiz-ru-1.docx"
    assert exported_file.media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert exported_file.content_bytes.startswith(b"PK")
    assert "Тренировочный квиз по географии" in document_text
    assert "Какой город является столицей России?" in document_text
    assert "Москва" in document_text
    assert "Казань" in document_text
    assert "Какая река протекает через Санкт-Петербург?" in document_text
    assert "Нева" in document_text
    assert "Ответы" in document_text
    assert "A. " in document_text


def test_default_export_registry_exposes_docx_exporter() -> None:
    exported_file = DEFAULT_QUIZ_EXPORT_REGISTRY.export(build_quiz(), "DOCX")
    document_text = read_docx_text(exported_file.content_bytes)

    assert "docx" in DEFAULT_QUIZ_EXPORT_REGISTRY.supported_formats()
    assert exported_file.filename == "quiz-ru-1.docx"
    assert "Тренировочный квиз по географии" in document_text
    assert "Ответы" in document_text


def test_docx_exporter_rejects_invalid_correct_option_index() -> None:
    quiz = build_quiz()
    invalid_question = replace(quiz.questions[0], correct_option_index=5)
    invalid_quiz = replace(quiz, questions=(invalid_question,))

    with pytest.raises(DomainValidationError, match="correct_option_index"):
        QuizDocxExporter().export(invalid_quiz)
